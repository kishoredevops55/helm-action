from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Helper function to set paragraph space after
def set_paragraph_spacing(paragraph, space_after=6):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(space_after)

# Helper function for line breaks
def add_line_break(doc, num_breaks=1):
    for _ in range(num_breaks):
        doc.add_paragraph()

# Create a new document
doc_reordered = Document()

# Title and Contact Information
title = doc_reordered.add_heading('KISHORE SHARMA', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(20)

subtitle = doc_reordered.add_paragraph('Senior Software Engineer')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.bold = True
subtitle_run.font.size = Pt(12)

contact_info = doc_reordered.add_paragraph(
    "Phone: (+91) 8015002126 | Email: kisaron5@gmail.com | Location: Chennai, India | LinkedIn: linkedin.com/in/kishore-sharma-sundaram-260a18120/"
)
contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_line_break(doc_reordered)

# Professional Summary Section
doc_reordered.add_heading('Professional Summary', level=1)
summary = doc_reordered.add_paragraph(
    "Highly skilled SRE and DevOps Engineer with over 7 years of experience in CI/CD automation, software configuration management, and infrastructure management. "
    "Expertise in implementing and supporting CI/CD pipelines, build & release management, and automation infrastructure for monitoring. Proficient in various DevOps "
    "and cloud platforms, with a focus on enhancing system reliability and operational efficiency."
)
set_paragraph_spacing(summary)

add_line_break(doc_reordered)

# Key Skills Section
doc_reordered.add_heading('Key Skills & Tools', level=1)

# Structure skills in two columns for a better look
skill_categories = {
    "Programming Languages": ["Python", "Shell Scripting", "Groovy"],
    "DevOps Tools": ["Jenkins", "GitLab CI", "GitHub Actions", "AutoCM", "Docker", "Kubernetes (Helm 3)", "Ansible", "Azure DevOps", "Maven"],
    "Version Control": ["GIT", "CA Harvest", "Subversion"],
    "Cloud Platforms": ["AWS", "Azure"],
    "Monitoring & Observability": ["Grafana Enterprise", "Grafana OSS", "Grafana Mimir (Metrics)", "Loki (Logs)", "Tempo (Traces)", "Prometheus", "Grafana Cloud Stack"],
    "Artifact Management": ["Nexus"],
    "Atlassian Tools": ["JIRA", "Confluence"],
    "Ticketing Tools": ["JIRA", "CA Service Desk", "BMC Remedy", "ServiceNow"],
    "Infrastructure as Code (IaC)": ["Terraform"],
}

# Add skills in a two-column format for better readability
skills_table = doc_reordered.add_table(rows=0, cols=2)
skills_table.style = 'Table Grid'

for category, skills in skill_categories.items():
    row_cells = skills_table.add_row().cells
    row_cells[0].paragraphs[0].add_run(f"{category}: ").bold = True
    row_cells[0].paragraphs[0].add_run(", ".join(skills))
    
# Adjusting column width
for row in skills_table.rows:
    row.cells[0].width = Pt(300)
    row.cells[1].width = Pt(300)

add_line_break(doc_reordered)

# Work Experience Section
doc_reordered.add_heading('Work Experience', level=1)

experiences_reordered = [
    {
        "position": "Sr Software Engineer - Cognizant",
        "dates": "11/2023 - Present",
        "responsibilities": [
            "Managed observability platforms for continuous monitoring, enhancing detection accuracy by 30%.",
            "Developed CI/CD pipelines with GitHub Actions and Jenkins, reducing deployment time by 50%.",
            "Managed microservices on AKS, improving resource utilization by 20%.",
            "Used Grafana Enterprise for metrics and tracing, increasing troubleshooting efficiency by 40%.",
            "Enhanced productivity with GitHub Copilot and VS Code."
        ]
    },
    # Additional experiences...
]

for job in experiences_reordered:
    experience_paragraph = doc_reordered.add_paragraph()
    experience_paragraph.add_run(f"{job['position']} ({job['dates']})").bold = True
    for responsibility in job['responsibilities']:
        doc_reordered.add_paragraph(f"   - {responsibility}", style='List Bullet')
    add_line_break(doc_reordered, num_breaks=1)

# Education Section
doc_reordered.add_heading('Education', level=1)
education_entries = [
    {"degree": "B.Tech in Computer Science", "institution": "Anna University", "year": "2015", "score": "8.5 CGPA"}
]

for edu in education_entries:
    p = doc_reordered.add_paragraph()
    p.add_run(f"{edu['degree']} - {edu['institution']} ({edu['year']})")
    if "score" in edu:
        p.add_run(f", Score: {edu['score']}")
        
add_line_break(doc_reordered)

# Projects Section
doc_reordered.add_heading('Projects', level=1)
projects = [
    "CI/CD pipeline setup for microservices architecture using Jenkins and Kubernetes.",
    "Infrastructure monitoring with Prometheus and Grafana.",
    "Automation of deployment processes using Ansible and Terraform."
]

for project in projects:
    doc_reordered.add_paragraph(f"- {project}", style='List Bullet')

# Languages Section
doc_reordered.add_heading('Languages', level=1)
doc_reordered.add_paragraph("Tamil, English")

# Save the document
output_path_reordered_docx = "Kishore_Resume_Final_Formatted.docx"
doc_reordered.save(output_path_reordered_docx)

output_path_reordered_docx
