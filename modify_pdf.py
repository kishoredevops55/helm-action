from docx import Document
import sys

def modify_document(input_path, output_path):
    # Open the document
    doc = Document(input_path)
    
    # Define replacements clearly
    replacements = {
        "Total Contribution\n1000.00": "Total Contribution\n40000.00",
        "Current Invested Amount\n1001.78": "Current Invested Amount\n40000.00",
        "Number of Contributions\n2": "Number of Contributions\n1",
        "Current Valuation\n995.41": "Current Valuation\n40000.00"
    }

    # Track replacements made
    changes_made = False

    # Iterate over paragraphs and apply replacements
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                print(f"Replacing in paragraph: '{old_text}' with '{new_text}'")
                para.text = para.text.replace(old_text, new_text)
                changes_made = True

    # Iterate over table cells for replacements
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    if old_text in cell.text:
                        print(f"Replacing in table cell: '{old_text}' with '{new_text}'")
                        cell.text = cell.text.replace(old_text, new_text)
                        changes_made = True

    # Save modified document if changes were made
    if changes_made:
        doc.save(output_path)
        print(f"Document saved as {output_path}")
    else:
        print("No matching text found; no changes made.")

# Run the modification function if called directly
if __name__ == "__main__":
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    modify_document(input_path, output_path)
